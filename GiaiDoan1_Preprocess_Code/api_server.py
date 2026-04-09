import re
import io
import torch
import numpy as np
import os
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from sentence_transformers import SentenceTransformer, util
from typing import List, Set, Dict, Optional
from dotenv import load_dotenv
from groq import Groq

# Load environment variables
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
groq_client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None

app = FastAPI(title="AI CV Matcher Pro", version="10.0")

# THÊM CORS ĐỂ FRONTEND GỌI ĐƯỢC API
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], # Cho phép tất cả (localhost:3000)
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

MODEL_PATH = "/home/hieu/Downloads/CV_MATCHING-main/GiaiDoan1_Preprocess_Code/models/e5_synthetic_model"
# ÉP CHẠY TRÊN CPU ĐỂ TIẾT KIỆM RAM (Tránh lag máy IdeaPad 8GB)
device = "cpu" 

try:
    model = SentenceTransformer(MODEL_PATH, device=device)
    print(f"✅ Loaded fine-tuned model on {device}")
except:
    model = SentenceTransformer("intfloat/multilingual-e5-small", device=device)
    print(f"✅ Fallback to E5-Small on {device}")

# =========================================================
# 2. DICTIONARY & LOGIC (GIỮ LẠI BẢN CŨ & TỐI ƯU)
# =========================================================
TECH_ALIASES = {
    "python": ["python", "py"], "java": ["java"], "javascript": ["javascript", "js", "jscript"],
    "typescript": ["typescript", "ts"], "react": ["react", "reactjs", "react.js"],
    "nodejs": ["nodejs", "node.js", "node"], "sql": ["sql", "mysql", "postgresql", "postgres"],
    "docker": ["docker", "containerization"], "aws": ["aws", "amazon web services"],
    "machine learning": ["machine learning", "ml", "ai"], "c#": ["c#", "csharp"],
    ".net": [".net", "dotnet", "asp.net"], "php": ["php", "laravel"]
}

CRITICAL_SKILLS = {"python", "java", "javascript", "react", "sql", "aws", "docker"}

# =========================================================
# 3. UTILS & PREPROCESSING
# =========================================================
def clean_text(text: str) -> str:
    text = text.lower()
    text = re.sub(r"\s+", " ", text) # Xóa xuống dòng, khoảng trắng thừa
    # Giữ lại các ký tự đặc biệt của ngành IT
    text = re.sub(r"[^a-z0-9\+\#\.\s]", " ", text)
    return text.strip()

def extract_keywords(text: str) -> Set[str]:
    text_lower = text.lower()
    # PREVENT FALSE POSITIVES: Ignore keywords inside negative contexts
    text_lower = re.sub(r"(chưa có|không có|không biết|chưa từng|điểm yếu|weakness)[^.]*", "", text_lower)
    
    found = set()
    cleaned = clean_text(text_lower)
    for canonical, aliases in TECH_ALIASES.items():
        for alias in aliases:
            # Regex boundary để tránh match nhầm 'java' trong 'javascript'
            pattern = rf"(?<!\w){re.escape(alias)}(?!\w)"
            if re.search(pattern, cleaned):
                found.add(canonical)
                break
    return found

# =========================================================
# 4. LONG TEXT HANDLING (CHIA ĐOẠN ĐỂ AI KHÔNG BỎ SÓT)
# =========================================================
def get_embedding(text: str, prefix: str = "passage: "):
    """Chia văn bản dài thành các đoạn nhỏ và lấy trung bình vector"""
    max_length = 400 # Token limit an toàn cho E5
    words = text.split()
    chunks = [" ".join(words[i:i + max_length]) for i in range(0, len(words), max_length)]
    
    if not chunks: 
        return torch.zeros((1, 384))
    
    # Thêm prefix cho mỗi chunk
    prefixed_chunks = [f"{prefix}{chunk}" for chunk in chunks]
    chunk_embeddings = model.encode(prefixed_chunks, convert_to_tensor=True, normalize_embeddings=True)
    
    # Trả về vector trung bình của các đoạn
    return torch.mean(chunk_embeddings, dim=0, keepdim=True)

# =========================================================
# 5. CORE MATCHING LOGIC
# =========================================================
class QAAnswer(BaseModel):
    question_id: str
    question_text: str
    answer: str
    type: str # "YES_NO" or "TEXT"

class MatchRequest(BaseModel):
    jd: str
    cv: str
    qa_answers: Optional[List[QAAnswer]] = None

class BatchCV(BaseModel):
    id: str # ID hoặc tên file
    text: str

class BatchMatchRequest(BaseModel):
    jd: str
    cvs: List[BatchCV]

class CVImprovementRequest(BaseModel):
    jd: str
    cv: str

class QuestionRequest(BaseModel):
    jd: str

@app.post("/api/generate-questions")
async def generate_questions(data: QuestionRequest):
    """Sử dụng Llama-3 để sinh câu hỏi sàng lọc từ JD"""
    if not groq_client:
        return {"error": "Groq API Key not configured"}
    
    try:
        prompt = f"""
        Dựa vào Job Description (JD) sau đây, hãy tạo tối đa 5 câu hỏi sàng lọc ứng viên.
        Yêu cầu:
        1. 2 câu hỏi dạng 'YES_NO' cho các yêu cầu bắt buộc (Hard Filter).
        2. 3 câu hỏi dạng 'TEXT' về kiến thức chuyên môn sâu (Soft Scoring).
        3. Các câu hỏi phải xoáy sâu vào các kỹ năng cốt lõi trong JD.
        4. Trả về định dạng JSON thuần túy (Array of Objects), không có văn bản giải thích nào khác.
        
        Định dạng mẫu:
        [
          {{"id": "q1", "text": "...", "type": "YES_NO"}},
          {{"id": "q2", "text": "...", "type": "TEXT"}}
        ]
        
        JD: {data.jd[:2000]}
        """
        
        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"} if "llama-3.1" in "llama-3.3-70b-versatile" else None
        )
        
        response_content = completion.choices[0].message.content
        # Trích xuất JSON nếu LLM trả lời kèm text
        import json
        try:
            questions = json.loads(response_content)
            # Nếu Groq trả về object có key "questions" hoặc tương tự
            if isinstance(questions, dict):
                for key in ["questions", "data", "result"]:
                    if key in questions:
                        questions = questions[key]
                        break
        except:
            # Fallback regex nếu JSON bị lỗi nhẹ
            match = re.search(r'\[\s*\{.*\}\s*\]', response_content, re.DOTALL)
            if match:
                questions = json.loads(match.group())
            else:
                raise Exception("Could not parse AI response as JSON")

        return {"success": True, "questions": questions[:5]}
    except Exception as e:
        print(f"Error generating questions: {e}")
        return {"error": str(e)}

@app.post("/api/batch-match")
async def batch_match_cv_jd(data: BatchMatchRequest):
    """Xử lý HÀNG LOẠT CV bằng Batch Encoding (Tối ưu NCKH)"""
    try:
        jd_raw = data.jd
        if not jd_raw or not data.cvs:
            raise HTTPException(status_code=400, detail="Thiếu dữ liệu")

        # 1. Tiền xử lý & Trích xuất Keyword 1 lần cho JD
        jd_clean = clean_text(jd_raw)
        jd_keys = extract_keywords(jd_raw)
        jd_emb = get_embedding(jd_clean, prefix="query: ") # (1, 384)

        # 2. Chuẩn bị Batch dữ liệu CV
        cv_texts = []
        cv_metadatas = []
        for cv_item in data.cvs:
            cv_clean = clean_text(cv_item.text)
            cv_texts.append(f"passage: {cv_clean[:2000]}") # Tránh quá dài
            cv_metadatas.append({
                "id": cv_item.id,
                "keys": extract_keywords(cv_item.text)
            })

        # 3. BATCH ENCODING (Đây là phần tối ưu chính)
        # Thay vì loop, ta đẩy cả List vào model
        # 3. BATCH ENCODING (Tối ưu: thêm batch_size để không gây spike RAM quá lớn)
        print(f"⚡ Batch Processing {len(cv_texts)} CVs...")
        all_cv_embs = model.encode(
            cv_texts, 
            batch_size=8, 
            convert_to_tensor=True, 
            normalize_embeddings=True,
            show_progress_bar=True
        ) # (N, 384)

        # 4. Tính toán Similarity hàng loạt bằng Ma trận
        # Cosine Similarity của vector đã normalize = Dot Product
        # jd_emb: (1, 384), all_cv_embs: (N, 384) -> Kết quả: (1, N)
        cos_sims = util.cos_sim(jd_emb, all_cv_embs)[0] 

        results = []
        for i, sim_score in enumerate(cos_sims):
            raw_sim = sim_score.item()
            cv_meta = cv_metadatas[i]
            
            # Tính điểm Keyword
            matched = jd_keys.intersection(cv_meta['keys'])
            missing = jd_keys.difference(cv_meta['keys'])
            
            # Chuẩn hóa Semantic Calibration (0.75 - 0.95 -> 0-100 cho model Fine-tuned)
            norm_semantic = max(0.0, min(1.0, (raw_sim - 0.75) / 0.20))
            kw_score = len(matched) / len(jd_keys) if jd_keys else norm_semantic
            
            final_score = (norm_semantic * 0.70 + kw_score * 0.30) * 100
            
            critical_missing = CRITICAL_SKILLS.intersection(missing)
            final_score -= (len(critical_missing) * 8)
            if len(jd_keys) > 0 and len(matched) == 0:
                final_score *= 0.5
            
            final_score = max(5.0, min(99.5, round(final_score, 1)))

            results.append({
                "id": cv_meta['id'],
                "score": final_score,
                "status": "EXCELLENT" if final_score >= 75 else "POTENTIAL" if final_score >= 60 else "CONSIDER" if final_score >= 40 else "REJECTED",
                "matched_count": len(matched),
                "semantic_sim": round(raw_sim, 4)
            })

        results.sort(key=lambda x: x['score'], reverse=True)
        return {"success": True, "count": len(results), "leaderboard": results}

    except Exception as e:
        import traceback
        print(traceback.format_exc())
        return {"error": str(e)}

@app.post("/api/match")
async def match_cv_jd(data: MatchRequest):
    try:
        jd_raw, cv_raw = data.jd, data.cv
        if not jd_raw or not cv_raw:
            raise HTTPException(status_code=400, detail="Dữ liệu không được để trống")

        # 1. Trích xuất Keywords (Thực tế)
        jd_keys = extract_keywords(jd_raw)
        cv_keys = extract_keywords(cv_raw)
        matched = jd_keys.intersection(cv_keys)
        missing = jd_keys.difference(cv_keys)

        # 2. Tính Semantic Similarity (AI)
        jd_emb = get_embedding(clean_text(jd_raw), prefix="query: ")
        cv_emb = get_embedding(clean_text(cv_raw), prefix="passage: ")
        
        raw_sim = util.cos_sim(jd_emb, cv_emb).item()

        # 3. Chuẩn hóa điểm Semantic Calibration (0.75 - 0.95 -> 0 - 100)
        norm_semantic = (raw_sim - 0.75) / 0.20
        norm_semantic = max(0.0, min(1.0, norm_semantic))

        # 4. Tính điểm Keyword
        kw_score = len(matched) / len(jd_keys) if jd_keys else norm_semantic

        # 5. [NEW] Chấm điểm Q&A (Hybrid Screening)
        qa_score = 1.0 # Mặc định 100% nếu không có Q&A
        qa_penalty = 0
        qa_critique = ""
        
        if data.qa_answers:
            correct_qa = 0
            total_qa = len(data.qa_answers)
            for qa in data.qa_answers:
                if qa.type == "YES_NO":
                    ans_clean = qa.answer.strip().lower()
                    if ans_clean == "" or ans_clean in ["no", "không", "n"]:
                        qa_penalty += 15 # Trừ 15 điểm mỗi câu Hard Filter "No" hoặc bỏ trống
                else:
                    # Dùng E5 đo độ khớp câu trả lời với JD
                    ans_emb = get_embedding(clean_text(qa.answer), prefix="passage: ")
                    # So sánh với JD (đã có jd_emb)
                    ans_sim = util.cos_sim(ans_emb, jd_emb).item()
                    # Normalize sim của câu trả lời (thường thấp hơn CV nên offset khác)
                    norm_ans = max(0.0, min(1.0, (ans_sim - 0.70) / 0.20))
                    correct_qa += norm_ans
            
            # Tính trung bình điểm Q&A (chỉ tính cho phần TEXT)
            text_q_count = len([q for q in data.qa_answers if q.type == "TEXT"])
            if text_q_count > 0:
                qa_score = correct_qa / text_q_count
            
            qa_critique = f" (Xác minh Q&A: {qa_score*100:.1f}%)"

        # 6. Công thức Hybrid Nâng cao (AI 50% + Keywords 20% + QA 30%)
        # Nếu không có QA thì giữ tỷ lệ cũ (70-30)
        if data.qa_answers:
            final_score = (norm_semantic * 0.50 + kw_score * 0.20 + qa_score * 0.30) * 100
            final_score -= qa_penalty
        else:
            final_score = (norm_semantic * 0.70 + kw_score * 0.30) * 100

        # 7. Penalty & Logic bổ trợ
        critical_missing = CRITICAL_SKILLS.intersection(missing)
        final_score -= (len(critical_missing) * 8) 

        if len(jd_keys) > 0 and len(matched) == 0:
            final_score *= 0.5 

        final_score = max(5.0, min(99.8, round(final_score, 1)))

        # 8. Phân loại ứng viên
        status = "REJECTED"
        if final_score >= 75: status = "EXCELLENT"
        elif final_score >= 60: status = "POTENTIAL"
        elif final_score >= 40: status = "CONSIDER"

        # 9. Xây dựng bản nhận xét chi tiết (Detailed Critique)
        critique = []
        sim_percent = norm_semantic * 100

        if final_score >= 85:
            critique.append(f"🌟 **Ứng viên xuất sắc.** Điểm: **{final_score:.1f}%** (AI CV: {sim_percent:.1f}%{qa_critique}).")

            if not missing:
                critique.append("Hồ sơ hoàn hảo, đáp ứng đầy đủ mọi yêu cầu kỹ thuật.")
            else:
                critique.append(f"Khớp {len(matched)}/{len(jd_keys)} kỹ năng: {', '.join(sorted(matched))}.")
                critique.append(f"Lưu ý nhỏ: Cần kiểm tra thêm về {', '.join(sorted(missing))}.")

        elif final_score >= 60:
            critique.append(f"✅ **Ứng viên tiềm năng.** Điểm tổng hợp đạt **{final_score:.1f}%** (Độ khớp AI ngữ nghĩa: {sim_percent:.1f}%).")

            if matched:
                critique.append(f"Điểm mạnh: Sở hữu các kỹ năng quan trọng như {', '.join(sorted(matched))}.")

            if critical_missing:
                critique.append(f"⚠️ **Rào cản:** Thiếu một số công nghệ lõi: {', '.join(sorted(critical_missing))}.")
            elif missing:
                critique.append(f"Cần bồi dưỡng thêm về: {', '.join(sorted(missing))}.")

        elif final_score >= 40:
            critique.append(f"🧐 **Ứng viên cần cân nhắc thêm.** Điểm tổng hợp đạt **{final_score:.1f}%** (Độ khớp AI ngữ nghĩa: {sim_percent:.1f}%).")

            if matched:
                critique.append(f"Có nền tảng về {', '.join(sorted(matched))}, nhưng chưa đủ bao quát JD.")

            if critical_missing:
                critique.append(f"🚩 **Cảnh báo:** Thiếu hụt nghiêm trọng các kỹ năng cốt lõi: {', '.join(sorted(critical_missing))}.")

        else:
            critique.append(f"❌ **Chưa phù hợp.** Trọng tâm hồ sơ hoàn toàn lệch so với yêu cầu vị trí.")

            if critical_missing:
                critique.append(f"Hồ sơ thiếu các kỹ năng bắt buộc: {', '.join(sorted(critical_missing))}.")

            critique.append("Khuyến dùng: Tìm kiếm các ứng viên có nền tảng công nghệ sát thực tế hơn hoặc xem xét ứng viên này cho các vị trí hỗ trợ/thực tập.")

        full_explanation = " ".join(critique)


        return {
            "score": final_score,
            "status": status,
            "analysis": {
                "semantic_sim": round(raw_sim, 4),
                "matched_skills": sorted(list(matched)),
                "missing_skills": sorted(list(missing)),
                "critical_missing": sorted(list(critical_missing))
            },
            "explanation": full_explanation
        }

    except Exception as e:
        return {"error": str(e)}


# =========================================================
# CHỨC NĂNG MỚI: SỰ CỎN CƯ CHỈNH SỬA CV TỰ ĐỘNG
# =========================================================
@app.post("/api/suggest-cv-improvements")
async def suggest_cv_improvements(data: CVImprovementRequest):
    """
    Phân tích CV và gợi ý cách chỉnh sửa để khớp hơn với JD.
    
    Trả về:
    - missing_skills: Các kỹ năng cần thêm
    - suggestions: Danh sách gợi ý cụ thể để chỉnh sửa CV
    - rewritten_sections: Các đoạn CV được viết lại để phù hợp hơn
    """
    try:
        jd_raw, cv_raw = data.jd, data.cv
        if not jd_raw or not cv_raw:
            raise HTTPException(status_code=400, detail="Dữ liệu không được để trống")

        # 1. Trích xuất Keywords từ JD và CV
        jd_keys = extract_keywords(jd_raw)
        cv_keys = extract_keywords(cv_raw)
        matched = jd_keys.intersection(cv_keys)
        missing = jd_keys.difference(cv_keys)
        critical_missing = CRITICAL_SKILLS.intersection(missing)

        # 2. Phân tích JD để tìm yêu cầu chính
        jd_requirements = _extract_job_requirements(jd_raw, jd_keys)

        # 3. Tạo danh sách gợi ý cụ thể
        suggestions = []
        priority_mapping = {}  # Lưu trữ prioritize cho từng kỹ năng

        # Gợi ý 1: Thêm từ khóa công nghệ thiếu
        if missing:
            missing_sorted = sorted(list(missing))
            if critical_missing:
                suggestions.append({
                    "type": "ADD_CRITICAL_SKILLS",
                    "priority": "HIGH",
                    "title": "🚨 Thêm kỹ năng bắt buộc",
                    "description": f"JD bắt buộc yêu cầu: {', '.join(sorted(critical_missing))}. CV của bạn chưa nhắc đến các kỹ năng này.",
                    "action": f"Thêm vào phần 'Kỹ năng' hoặc mô tả công việc: {', '.join(critical_missing)}",
                    "example": f"Ví dụ: 'Kinh nghiệm với {critical_missing.pop()} trong các dự án thực tế'"
                })
                # Phục hồi critical_missing
                critical_missing = CRITICAL_SKILLS.intersection(missing)

            non_critical_missing = set(missing_sorted) - CRITICAL_SKILLS
            if non_critical_missing:
                suggestions.append({
                    "type": "ADD_OPTIONAL_SKILLS",
                    "priority": "MEDIUM",
                    "title": "💡 Bổ sung kỹ năng bổ trợ",
                    "description": f"JD nhắc đến: {', '.join(sorted(non_critical_missing))}. Nên thêm nếu bạn có kinh nghiệm.",
                    "action": f"Xem xét thêm: {', '.join(sorted(non_critical_missing))}",
                    "example": f"Ví dụ: 'Có kinh nghiệm cơ bản với {non_critical_missing.pop()}'"
                })

        # Gợi ý 2: Tối ưu hóa phần Summary/Objective
        job_titles_in_jd = _extract_job_titles(jd_raw)
        if job_titles_in_jd:
            suggestions.append({
                "type": "OPTIMIZE_SUMMARY",
                "priority": "HIGH",
                "title": "✒️ Viết lại Objective/Summary",
                "description": "Phần mở đầu CV nên phản ánh rõ lĩnh vực/vị trí mục tiêu.",
                "action": f"Thay đổi: 'Seeking {job_titles_in_jd[0]} with strong background in {', '.join(list(matched)[:3]) if matched else 'relevant'} technologies'",
                "example": f"Thay vì 'Looking for a good opportunity', hãy viết: '{job_titles_in_jd[0]} with expertise in {', '.join(list(matched)[:2]) if matched else 'technology implementation'}'"
            })

        # Gợi ý 3: Sắp xếp lại phần Experience
        if missing and matched:
            suggestions.append({
                "type": "PRIORITIZE_EXPERIENCE",
                "priority": "MEDIUM",
                "title": "👨‍💼 Sắp xếp lại phần kinh nghiệm",
                "description": f"Các công việc liên quan đến {', '.join(list(matched)[:2])} nên được đặt lên trên cùng.",
                "action": "Sắp xếp lại Experience section: những công việc sử dụng công nghệ trong JD lên trên.",
                "example": "Nếu có kinh nghiệm làm Backend, hãy đặt job đó trước job Frontend nếu JD tìm Backend"
            })

        # Gợi ý 4: Thêm mô tả chi tiết hơn cho các kỹ năng có
        if matched:
            suggestions.append({
                "type": "ENHANCE_SKILL_DESCRIPTION",
                "priority": "MEDIUM",
                "title": "📝 Mô tả chi tiết hơn các kỹ năng đã có",
                "description": f"CV có {len(matched)}/{len(jd_keys)} kỹ năng yêu cầu. Cần làm rõ hơn bằng cách thêm context.",
                "action": f"Cho mỗi công việc, thêm: 'Dùng {', '.join(list(matched)[:2])} để xây dựng...' hoặc 'Đạt giải thưởng/chỉ tiêu khi áp dụng {', '.join(list(matched)[:2])}'",
                "example": "Thay vì 'Used Python', hãy viết: 'Developed data pipeline using Python and SQL, processing 10M+ records daily'"
            })

        # Gợi ý 5: Bổ sung phần Projects nếu cần
        if critical_missing:
            suggestions.append({
                "type": "ADD_PROJECTS",
                "priority": "HIGH",
                "title": "🛠️ Bổ sung phần Projects",
                "description": f"Để bù đắp khoảng trống về {', '.join(list(critical_missing)[:2])}, hãy thêm phần 'Side Projects'.",
                "action": "Thêm section 'Projects' với các dự án free/personal project sử dụng công nghệ bắt buộc",
                "example": f"Ví dụ: 'Xây dựng {critical_missing.pop()}-based web app với [các kỹ năng]' (ngay cả hobby project cũng được)"
            })

        # 4. Tạo phần CV viết lại mẫu
        rewritten_summary = _generate_improved_cv_summary(
            cv_raw, jd_keys, matched, job_titles_in_jd
        )

        # Sắp xếp gợi ý theo priority
        high_priority = [s for s in suggestions if s.get("priority") == "HIGH"]
        medium_priority = [s for s in suggestions if s.get("priority") == "MEDIUM"]
        suggestions = high_priority + medium_priority

        return {
            "success": True,
            "analysis": {
                "matched_skills": sorted(list(matched)),
                "missing_skills": sorted(list(missing)),
                "critical_missing": sorted(list(critical_missing)),
                "match_percentage": round(len(matched) / len(jd_keys) * 100, 1) if jd_keys else 0
            },
            "suggestions": suggestions,
            "improved_summary": rewritten_summary,
            "action_steps": [
                f"1. Ưu tiên thêm: {', '.join(list(critical_missing)[:3]) if critical_missing else 'các kỹ năng bắt buộc'}",
                f"2. Viết lại Objective để phù hợp với '{job_titles_in_jd[0] if job_titles_in_jd else 'vị trí tìm kiếm'}'",
                f"3. Làm nổi bật {len(matched)} kỹ năng bạn đã có: {', '.join(list(matched)[:3])}",
                "4. Thêm examples/metrics cụ thể cho mỗi kỹ năng",
                "5. Bộ sơ cấu trúc lại: Experience → Projects → Skills"
            ]
        }

    except Exception as e:
        import traceback
        print(traceback.format_exc())
        return {"error": str(e)}


# =========================================================
# HỖ TRỢ: Các hàm phụ trợ cho chức năng suggestions
# =========================================================
def _extract_job_titles(text: str) -> List[str]:
    """Trích xuất các tiêu đề công việc phổ biến từ JD"""
    job_titles = []
    common_titles = [
        "software engineer", "developer", "backend", "frontend", "full-stack",
        "data scientist", "machine learning engineer", "devops", "cloud architect",
        "product manager", "project manager", "qa engineer", "data engineer"
    ]
    text_lower = text.lower()
    for title in common_titles:
        if title in text_lower:
            job_titles.append(title.title())
    return job_titles[:3] if job_titles else ["Software Professional"]


def _extract_job_requirements(text: str, skills: Set[str]) -> Dict[str, List[str]]:
    """Phân tích JD để tìm các yêu cầu chính"""
    requirements = {
        "core_skills": list(skills)[:5],
        "experience_years": [],
        "education": []
    }
    
    # Tìm kinh nghiệm năm
    years_pattern = r"(\d+)\s*(?:\+)?\s*years?"
    years_matches = re.findall(years_pattern, text.lower())
    if years_matches:
        requirements["experience_years"] = years_matches
    
    return requirements


def _generate_improved_cv_summary(cv_text: str, jd_keys: Set[str], 
                                   matched: Set[str], job_titles: List[str]) -> str:
    """
    Tạo phiên bản CV tóm tắt được cải thiện
    """
    title = job_titles[0] if job_titles else "Professional"
    matched_str = ", ".join(list(matched)[:3]) if matched else "relevant technologies"
    
    improved = f"""
    **[IMPROVED OBJECTIVE]**
    
    Experienced {title} with strong expertise in {matched_str}. 
    Proven track record in building scalable solutions and delivering results.
    Seeking to leverage skills in {', '.join(list(jd_keys)[:3])} to contribute to your team's success.
    
    **[KEY ENHANCEMENTS]**
    1. Specific skills mentioned: Focus on {matched_str}
    2. Measurable achievements: Add metrics (e.g., "Improved performance by X%", "Led team of Y people")
    3. Relevant projects: Highlight work using {', '.join(list(jd_keys)[:2])}
    
    **[SUGGESTED SKILLS SECTION]**
    Core: {', '.join(list(matched))}
    To Add: {', '.join(list(jd_keys - matched)[:3])}
    """
    
    return improved.strip()


# =========================================================
# 6. PDF / DOCX FILE PARSING ENDPOINT
# =========================================================
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Trích xuất text từ PDF – xử lý layout 2 cột thông minh"""
    try:
        import pdfplumber
        text_pages = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                width = page.width
                # Thử phát hiện layout 2 cột
                left_bbox  = (0,          0, width * 0.52, page.height)
                right_bbox = (width * 0.48, 0, width,      page.height)

                left_text  = page.within_bbox(left_bbox).extract_text()  or ""
                right_text = page.within_bbox(right_bbox).extract_text() or ""

                # Nếu cả 2 cột đều có nội dung → layout 2 cột
                if left_text.strip() and right_text.strip():
                    combined = left_text.strip() + "\n" + right_text.strip()
                else:
                    # Layout 1 cột bình thường
                    combined = page.extract_text() or ""

                text_pages.append(combined)

        return "\n".join(text_pages).strip()
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="Thư viện 'pdfplumber' chưa được cài. Chạy: pip install pdfplumber"
        )


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Trích xuất text từ file DOCX"""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Bao gồm cả text trong bảng (Skills table thường nằm trong table)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs).strip()
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="Thư viện 'python-docx' chưa được cài. Chạy: pip install python-docx"
        )


@app.post("/api/parse-file")
async def parse_file(file: UploadFile = File(...)):
    """
    Nhận file PDF / DOCX / TXT và trả về text đã trích xuất.
    Frontend dùng text này để điền vào textarea CV hoặc JD.
    """
    try:
        filename  = file.filename or ""
        ext       = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        file_bytes = await file.read()

        if not file_bytes:
            raise HTTPException(status_code=400, detail="File rỗng, không thể đọc.")

        if ext == "pdf":
            text = extract_text_from_pdf(file_bytes)
        elif ext in ("docx", "doc"):
            text = extract_text_from_docx(file_bytes)
        elif ext == "txt":
            text = file_bytes.decode("utf-8", errors="replace")
        else:
            raise HTTPException(
                status_code=415,
                detail=f"Định dạng '.{ext}' không được hỗ trợ. Chỉ chấp nhận: PDF, DOCX, TXT."
            )

        if not text.strip():
            raise HTTPException(
                status_code=422,
                detail="Không trích xuất được text từ file. File có thể là ảnh (scanned PDF) hoặc bị bảo vệ."
            )

        # Đếm số từ để validate
        word_count = len(text.split())

        return {
            "success": True,
            "filename": filename,
            "text": text,
            "word_count": word_count,
            "char_count": len(text)
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Lỗi xử lý file: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)
    